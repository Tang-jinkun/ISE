from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Agent/test/fixtures/north-sea-evacuation-interception.docx"

SCENARIO_PARAGRAPHS = [
    "06:00 - Blue Task Group operates one Boeing E-3A Sentry AWACS from coordinates:-5.800,57.400 to coordinates:-3.200,58.100 for continuous early warning patrol.",
    "06:05 - Blue Task Group launches a formation of four Rafale aircraft from coordinates:-4.900,56.900 to coordinates:-0.200,58.000. The formation keeps a stable lead aircraft.",
    "06:06 - Red Air Group launches a formation of four J-10 aircraft from coordinates:2.600,59.100 to coordinates:0.300,58.000. The formation keeps a stable lead aircraft.",
    "06:10 - The Boeing E-3A Sentry AWACS sends target information by data link to all four Blue Rafale aircraft.",
    "06:14 - The lead Blue Rafale launches one PL-15E missile from coordinates:-1.500,57.900 to coordinates:0.300,58.000 at the lead Red J-10.",
    "06:16 - The first PL-15E and the targeted J-10 reach coordinates:0.300,58.000 together. Their geometry intersects and the targeted J-10 is confirmed destroyed.",
    "06:16 - The lead Rafale sends terminal-guidance data by data link to the first PL-15E until impact.",
    "06:20 - A surviving Red J-10 launches one PL-15E missile from coordinates:0.600,58.100 to coordinates:-0.700,58.350 toward the Blue Rafale formation.",
    "06:22 - The second PL-15E ends at coordinates:-0.700,58.350 while the Rafale formation remains on its documented route. No geometric intersection is established and the outcome remains unconfirmed.",
    "06:22 - The Red J-10 sends weapon-status data by data link to the second PL-15E. No hit, destruction, or miss may be claimed.",
]


def set_run_font(run, *, size=None, color=None, bold=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_properties = table._tbl.tblPr
    table_width = table_properties.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    table_indent = table_properties.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_properties.append(table_indent)
    table_indent.set(qn("w:w"), "120")
    table_indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            properties = cell._tc.get_or_add_tcPr()
            cell_width = properties.first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                properties.append(cell_width)
            cell_width.set(qn("w:w"), str(widths[index]))
            cell_width.set(qn("w:type"), "dxa")


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("ISE | Grounded start/end scenario brief")
    set_run_font(header_run, size=9, color="646464")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Document-to-scene compilation fixture")
    set_run_font(footer_run, size=9, color="646464")


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.add_run(text)
    return paragraph


def add_metadata_table(doc):
    metadata = [
        ("Purpose", "Validate a self-contained North Sea interception report through DOCX evidence parsing."),
        ("Geometry policy", "Every moving actor statement records explicit start and end coordinate anchors."),
        ("Interaction policy", "Only the documented geometric intersection is destroyed; the second engagement remains unresolved."),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in metadata:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], "F2F4F7")
        for cell in cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    set_run_font(run, size=10)
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
    set_table_geometry(table, [2700, 6660])


def main():
    output = Path(OUT)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("NORTH SEA EVACUATION CORRIDOR INTERCEPTION")
    set_run_font(title_run, size=23, color="0B2545", bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run("Grounded start/end scenario brief for ISE Agent compilation")
    set_run_font(subtitle_run, size=14, color="373737")

    add_metadata_table(doc)
    doc.add_paragraph()

    doc.add_heading("1. Early-warning patrol and data-link picture", level=1)
    add_body(doc, SCENARIO_PARAGRAPHS[0])
    add_body(doc, SCENARIO_PARAGRAPHS[3])

    doc.add_heading("2. Fighter formations and grounded movement", level=1)
    add_body(doc, SCENARIO_PARAGRAPHS[1])
    add_body(doc, SCENARIO_PARAGRAPHS[2])

    doc.add_heading("3. Confirmed launch and destruction", level=1)
    for paragraph in SCENARIO_PARAGRAPHS[4:7]:
        add_body(doc, paragraph)

    doc.add_heading("4. Second approach remains unresolved", level=1)
    for paragraph in SCENARIO_PARAGRAPHS[7:]:
        add_body(doc, paragraph)

    doc.save(output)


if __name__ == "__main__":
    main()
