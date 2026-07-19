from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "agent/test/fixtures/start-end-indo-pak-interception.docx"


def set_cell_shading(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    props.append(shade)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    props = cell._tc.get_or_add_tcPr()
    margins = props.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        props.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    header = section.header.paragraphs[0]
    header.text = "ISE | Grounded start/end scenario brief"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.name = "Calibri"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    footer = section.footer.paragraphs[0]
    footer.text = "Document-to-scene compilation fixture"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.name = "Calibri"
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(100, 100, 100)


def add_body(doc, text, *, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def main():
    doc = Document()
    configure_document(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("INDO-PAK AIR INTERCEPTION")
    run.font.name = "Calibri"
    run.font.size = Pt(23)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("Grounded start/end scenario brief for ISE Agent compilation")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(55, 55, 55)

    metadata = [
        ("Purpose", "Validate DOCX -> EvidenceIR -> EventPlan -> runtime route synthesis"),
        ("Source geometry", "Every named moving actor below has an explicit start and end coordinate"),
        ("Interaction policy", "Only the confirmed geometric impact is destroyed; the second approach remains unresolved"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in metadata:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value
        set_cell_shading(row[0], "F2F4F7")
        for cell in row:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for r in paragraph.runs:
                    r.font.name = "Calibri"
                    r.font.size = Pt(10)
            for r in cell.paragraphs[0].runs:
                if cell is row[0]:
                    r.bold = True
    set_table_geometry(table, [2700, 6660])
    doc.add_paragraph()

    doc.add_heading("1. Early-warning patrol and data-link picture", level=1)
    add_body(doc, "Scenario geography: the source report names Adampur, Ambala, Minhas, and Rafiki as the operating areas. The following coordinate pairs are the authoritative route anchors; they intentionally do not depend on a pre-authored trajectory JSON.")
    add_body(doc, "08:00 - India's Netra AEW&CS, one aircraft, patrols from Adampur coordinates:75.660,31.430 to coordinates:76.300,31.600. The route is a continuous surveillance patrol.")
    add_body(doc, "08:00 - Pakistan's ZDK-03 AWACS, one aircraft, patrols from Minhas coordinates:73.100,33.870 to coordinates:72.800,33.400. This is a separate continuous surveillance patrol.")
    add_body(doc, "08:08 - Netra AEW&CS sends a live data link to the Indian Su-30MKI formation and the Indian Rafale formation. The link carries target information and guidance; it does not move either aircraft by itself.")

    doc.add_heading("2. Fighter formations and grounded movement", level=1)
    add_body(doc, "08:12 - India launches a formation of four Su-30MKI aircraft from coordinates:75.660,31.430 to coordinates:75.980,32.000. The four aircraft remain a single formation with a stable lead aircraft.")
    add_body(doc, "08:13 - India launches a formation of four Rafale aircraft from coordinates:76.820,30.370 to coordinates:76.500,31.800. The four aircraft remain a separate formation.")
    add_body(doc, "08:14 - Pakistan launches a formation of four JF-17 aircraft from coordinates:73.100,33.870 to coordinates:75.050,32.050. The lead JF-17 is the named target in the confirmed engagement below.")

    doc.add_heading("3. Confirmed launch and destruction", level=1)
    add_body(doc, "08:20 - One Indian Su-30MKI launches one PL-15E missile from coordinates:76.100,31.850 to coordinates:75.050,32.050 toward the Pakistani JF-17 lead aircraft.")
    add_body(doc, "08:22 - The PL-15E missile reaches coordinates:75.050,32.050 at the same time as the targeted JF-17. The missile intersects the aircraft, and the JF-17 is confirmed destroyed.")
    add_body(doc, "08:22 - The Su-30MKI sends a weapon-status data link to the launched PL-15E until terminal guidance is complete.")

    doc.add_heading("4. Second approach remains unresolved", level=1)
    add_body(doc, "08:25 - A Pakistani JF-17 launches a second PL-15E missile from coordinates:74.900,32.000 to coordinates:75.700,31.900 toward the Indian Rafale formation.")
    add_body(doc, "08:27 - The second approach is observed but no geometric intersection is established in the source report. The outcome is unresolved: do not claim a hit, destruction, or miss. The JF-17 continues its route and the Rafale formation remains present.")
    add_body(doc, "08:27 - The Pakistani ZDK-03 AWACS sends target information to the JF-17 formation, and the JF-17 sends a weapon-status data link to the second PL-15E.")

    doc.add_heading("5. Compilation notes", level=1)
    add_body(doc, "The coordinate pairs are the authoritative movement anchors for this fixture. No trajectory JSON is attached. When no catalog route matches these exact endpoints, the compiler must synthesize deterministic great-circle or intercept trajectories and preserve the evidence references.")
    add_body(doc, "Quantities written as four aircraft are exact formation quantities. The single aircraft and single missile statements are event-scoped actions and must not collapse the persistent formations. The confirmed destruction is valid only when the runtime geometry supports the stated interaction; the second approach must remain unresolved otherwise.")

    doc.save(OUT)


if __name__ == "__main__":
    main()
